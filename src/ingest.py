"""Document ingestion: load PDF/DOCX/TXT/MD -> text with metadata.

Extract text from each document and attach metadata (filename, page) so that
citations can later show "which file, which page".
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from . import config


@dataclass
class Document:
    """One unit of source text (a page for PDFs, whole file for others)."""

    text: str
    metadata: dict = field(default_factory=dict)


def _load_pdf(path: Path) -> list[Document]:
    """Extract per-page text AND tables via PyMuPDF.

    Tables are converted to Markdown and appended as their own block so chunking
    keeps them intact (a plain text dump would scramble rows/columns). Falls back
    to pypdf if PyMuPDF is unavailable.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return _load_pdf_pypdf(path)

    docs: list[Document] = []
    pdf = fitz.open(str(path))
    try:
        for page_num, page in enumerate(pdf, start=1):
            blocks: list[str] = []
            text = (page.get_text() or "").strip()
            if text:
                blocks.append(text)
            # structured tables -> Markdown (kept as atomic blocks)
            try:
                for tbl in page.find_tables().tables:
                    md = (tbl.to_markdown() or "").strip()
                    if md:
                        blocks.append("[Table]\n" + md)
            except Exception:  # noqa: BLE001 — table detection is best-effort
                pass
            combined = "\n\n".join(blocks).strip()
            if combined:
                docs.append(
                    Document(text=combined, metadata={"source": path.name, "page": page_num})
                )
    finally:
        pdf.close()
    return docs


def _load_pdf_pypdf(path: Path) -> list[Document]:
    """Fallback PDF loader (text only) when PyMuPDF isn't installed."""
    from pypdf import PdfReader

    docs: list[Document] = []
    reader = PdfReader(str(path))
    for page_num, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            docs.append(
                Document(text=text, metadata={"source": path.name, "page": page_num})
            )
    return docs


def _load_docx(path: Path) -> list[Document]:
    import docx  # python-docx

    document = docx.Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    if not text.strip():
        return []
    return [Document(text=text, metadata={"source": path.name, "page": 1})]


def _load_text(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    if not text.strip():
        return []
    return [Document(text=text, metadata={"source": path.name, "page": 1})]


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_text,
    ".md": _load_text,
}


def load_file(path: Path) -> list[Document]:
    """Load a single file into Document(s). Unsupported -> []."""
    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        return []
    try:
        return loader(path)
    except Exception as exc:  # noqa: BLE001 — one bad file shouldn't kill the batch
        print(f"[ingest] skipped {path.name}: {exc}")
        return []


def load_dir(data_dir: Path | str = config.DATA_DIR) -> list[Document]:
    """Load every supported document in a directory (recursively)."""
    data_dir = Path(data_dir)
    docs: list[Document] = []
    if not data_dir.exists():
        return docs
    for path in sorted(data_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in config.SUPPORTED_EXTS:
            docs.extend(load_file(path))
    return docs


def iter_supported(paths: Iterable[Path]) -> list[Document]:
    """Load a specific list of files (used by the UI after upload)."""
    docs: list[Document] = []
    for p in paths:
        docs.extend(load_file(Path(p)))
    return docs
